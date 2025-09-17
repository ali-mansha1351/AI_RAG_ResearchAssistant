"""Document processing utilities for various file types."""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
from dataclasses import dataclass

# PDF processing
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pymupdf4llm
    PYMUPDF4LLM_AVAILABLE = True
except ImportError:
    PYMUPDF4LLM_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Document processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

from config.settings import settings

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Container for processed document data."""
    content: str
    metadata: Dict[str, Any]
    chunks: Optional[List[LangChainDocument]] = None


class DocumentProcessor:
    """Processes various document types into text chunks."""
    
    def __init__(self, chunk_size: Optional[int] = None, chunk_overlap: Optional[int] = None):
        """Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size or settings.chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", ".", "!", "?", " ", ""]
        )
    
    def process_file(self, file_path: str, show_progress: bool = True) -> ProcessedDocument:
        """Process a file based on its extension with optimized performance.
        
        Args:
            file_path: Path to the file
            show_progress: Whether to show progress information
            
        Returns:
            ProcessedDocument containing text and metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = path.suffix.lower()
        file_size = path.stat().st_size
        
        if show_progress:
            logger.info(f"Processing {file_extension} file: {path.name} ({file_size / 1024 / 1024:.2f} MB)")
        
        # Extract text based on file type
        if file_extension == '.pdf':
            content, metadata = self._process_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            content, metadata = self._process_docx(file_path)
        elif file_extension == '.txt':
            content, metadata = self._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        if show_progress:
            logger.info(f"Extracted {len(content)} characters from {path.name}")
        
        # Create chunks
        chunks = self._create_chunks(content, metadata)
        
        if show_progress:
            logger.info(f"Created {len(chunks)} chunks from {path.name}")
        
        return ProcessedDocument(
            content=content,
            metadata=metadata,
            chunks=chunks
        )
    
    def _process_pdf(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process PDF file with version-compatible extraction."""
        content = ""
        metadata = {
            "source": file_path,
            "type": "pdf",
            "total_pages": 0
        }
        
        try:
            # Use PyMuPDF (fitz) with enhanced extraction
            doc = fitz.open(file_path)
            metadata["total_pages"] = len(doc)
            
            for page_num in range(len(doc)):
                try:
                    page = doc.load_page(page_num)
                    page_text = ""
                    
                    # Method 1: Basic text extraction with version compatibility
                    try:
                        # Try different API versions
                        try:
                            page_text = page.get_text()
                        except TypeError:
                            # Older versions might not have parameters
                            page_text = page.getText()
                    except AttributeError:
                        # Really old versions
                        try:
                            page_text = page.extractText()
                        except AttributeError:
                            page_text = ""
                    
                    # Method 2: Try extracting text in different ways if basic approach yields little
                    if len(page_text.strip()) < 100:
                        try:
                            # Try to get word-level text
                            words = page.get_text("words")
                            if words and isinstance(words, list):
                                words_text = " ".join([w[4] for w in words if len(w) > 4])
                                if len(words_text) > len(page_text):
                                    page_text = words_text
                        except (AttributeError, TypeError, IndexError):
                            pass
                    
                    # Method 3: Manual table-like content extraction without find_tables()
                    try:
                        # Simple grid detection that works across versions
                        rect = page.rect
                        # Divide page into a grid and extract text from each cell
                        cols, rows = 3, 5  # Simple 3x5 grid
                        width, height = rect.width / cols, rect.height / rows
                        
                        grid_text = "\nStructured Content:\n"
                        for row in range(rows):
                            row_text = []
                            for col in range(cols):
                                # Create a rectangle for this grid cell
                                cell_rect = fitz.Rect(
                                    rect.x0 + col * width,
                                    rect.y0 + row * height,
                                    rect.x0 + (col + 1) * width,
                                    rect.y0 + (row + 1) * height
                                )
                                # Extract text from this region
                                try:
                                    cell_text = page.get_text("text", clip=cell_rect)
                                except (TypeError, AttributeError):
                                    try:
                                        cell_text = page.getText("text", clip=cell_rect)
                                    except (TypeError, AttributeError):
                                        cell_text = ""
                                
                                if cell_text.strip():
                                    row_text.append(cell_text.strip())
                            
                            if row_text:
                                grid_text += " | ".join(row_text) + "\n"
                        
                        if len(grid_text) > 30:  # Only add if meaningful content was found
                            page_text += grid_text
                    except Exception as grid_error:
                        logger.debug(f"Grid extraction failed: {grid_error}")
                    
                    content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    
                    # Enhanced text extraction for structured content
                    if any(keyword in page_text.lower() for keyword in ['qr code', 'fields', 'data elements', 'mandatory', 'required']):
                        try:
                            # Extract text with more detail for pages containing QR code information
                            detailed_text = page.get_text("dict")
                            if isinstance(detailed_text, dict) and 'blocks' in detailed_text:
                                structured_content = "\nDetailed Content:\n"
                                for block in detailed_text['blocks']:
                                    if 'lines' in block:
                                        for line in block['lines']:
                                            if 'spans' in line:
                                                line_text = " ".join([span.get('text', '') for span in line['spans']])
                                                if line_text.strip():
                                                    structured_content += line_text + "\n"
                                content += structured_content
                        except Exception as detail_error:
                            logger.debug(f"Detailed extraction failed: {detail_error}")
                    
                except Exception as page_error:
                    logger.warning(f"Error processing page {page_num} in {file_path}: {page_error}")
                    content += f"\n--- Page {page_num + 1} [Extraction Error] ---\n"
            
            doc.close()
            logger.info(f"Successfully processed PDF with PyMuPDF: {file_path}")
            
        except Exception as e:
            logger.error(f"PyMuPDF processing failed for {file_path}: {e}")
            # Fall back to PyPDF2
            content = self._process_pdf_pypdf2(file_path)
        
        content = self._clean_text(content)
        return content, metadata
    
    def _process_pdf_pypdf2(self, file_path: str) -> str:
        """Process PDF using PyPDF2."""
        content = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                content += page.extract_text()
        
        logger.info(f"Processed PDF with PyPDF2: {file_path}")
        return content
    
    def _process_docx(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        doc = Document(file_path)
        content = ""
        
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        
        metadata = {
            "source": file_path,
            "type": "docx",
            "paragraphs": len(doc.paragraphs)
        }
        
        content = self._clean_text(content)
        logger.info(f"Processed DOCX: {file_path}")
        
        return content, metadata
    
    def _process_txt(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()
        
        metadata = {
            "source": file_path,
            "type": "txt",
            "size": len(content)
        }
        
        content = self._clean_text(content)
        logger.info(f"Processed TXT: {file_path}")
        
        return content, metadata
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Skip empty lines
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _create_chunks(self, content: str, metadata: Dict[str, Any]) -> List[LangChainDocument]:
        """Split content into chunks with enhanced handling for structured data."""
        # Pre-process content to identify and preserve important sections
        enhanced_content = self._enhance_content_for_chunking(content)
        
        chunks = self.text_splitter.split_text(enhanced_content)
        
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata["chunk_id"] = i
            chunk_metadata["chunk_size"] = len(chunk)
            
            # Add semantic tags for better retrieval
            if any(keyword in chunk.lower() for keyword in ['qr code', 'mandatory', 'required fields', 'data elements']):
                chunk_metadata["content_type"] = "structured_data"
                chunk_metadata["importance"] = "high"
            
            documents.append(LangChainDocument(
                page_content=chunk,
                metadata=chunk_metadata
            ))
        
        logger.info(f"Created {len(documents)} chunks from {metadata['source']}")
        return documents
    
    def _enhance_content_for_chunking(self, content: str) -> str:
        """Enhance content to better preserve important sections during chunking."""
        lines = content.split('\n')
        enhanced_lines = []
        
        for i, line in enumerate(lines):
            # Add section markers for important content
            if any(keyword in line.lower() for keyword in ['qr code', 'mandatory', 'required', 'data elements']):
                # Add emphasis to ensure this content stays together
                enhanced_lines.append(f"\n=== IMPORTANT SECTION ===")
                enhanced_lines.append(line)
                
                # Include context around important lines
                context_start = max(0, i-2)
                context_end = min(len(lines), i+3)
                for j in range(context_start, context_end):
                    if j != i and j < len(lines):
                        enhanced_lines.append(lines[j])
                enhanced_lines.append("=== END SECTION ===\n")
            else:
                enhanced_lines.append(line)
        
        return '\n'.join(enhanced_lines)
    
    def process_text(self, text: str, source: str = "text_input") -> ProcessedDocument:
        """Process raw text string."""
        metadata = {
            "source": source,
            "type": "text",
            "size": len(text)
        }
        
        cleaned_text = self._clean_text(text)
        chunks = self._create_chunks(cleaned_text, metadata)
        
        return ProcessedDocument(
            content=cleaned_text,
            metadata=metadata,
            chunks=chunks
        )